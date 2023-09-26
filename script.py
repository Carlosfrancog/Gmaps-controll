import os
import smtplib
from datetime import datetime
import email.message
import tkinter as tk
from tkinter import messagebox
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx import Document


def verify():
    # new window
    verify_window = tk.Toplevel(root)
    verify_window.title(" verify")

    # form labels
    operador_label = tk.Label(verify_window, text=f"Operador: {operador_entry.get()}")
    operador_label.pack()

    nome_fantasia_label = tk.Label(verify_window, text=f"Nome Fantasia: {nome_fantasia_entry.get()}")
    nome_fantasia_label.pack()

    razao_social_label = tk.Label(verify_window, text=f"Razão Social: {razao_social_entry.get()}")
    razao_social_label.pack()

    cnpj_label = tk.Label(verify_window, text=f"CNPJ/CPF: {cnpj_entry.get()}")
    cnpj_label.pack()

    ramo_atividade_label = tk.Label(verify_window, text=f"Ramo de Atividade: {ramo_atividade_entry.get()}")
    ramo_atividade_label.pack()

    endereco_label = tk.Label(verify_window, text=f"Endereço: {endereco_entry.get()}")
    endereco_label.pack()

    cidade_label = tk.Label(verify_window, text=f"Cidade: {cidade_entry.get()}")
    cidade_label.pack()

    cep_label = tk.Label(verify_window, text=f"CEP: {cep_entry.get()}")
    cep_label.pack()

    telefone_label = tk.Label(verify_window, text=f"Telefone: {telefone_entry.get()}")
    telefone_label.pack()

    autorizante_label = tk.Label(verify_window, text=f"Autorizante: {autorizante_entry.get()}")
    autorizante_label.pack()

    cargo_label = tk.Label(verify_window, text=f"Cargo: {cargo_entry.get()}")
    cargo_label.pack()

    data_venda_label = tk.Label(verify_window, text=f"Data Venda: {data_venda_entry.get()}")
    data_venda_label.pack()

    vencimento_label = tk.Label(verify_window, text=f"Vencimento: {vencimento_entry.get()}")
    vencimento_label.pack()

    valor_label = tk.Label(verify_window, text=f"Valor: {valor_entry.get()}")
    valor_label.pack()

    email_label = tk.Label(verify_window, text=f"E-mail: {email_entry.get()}")
    email_label.pack()
    
    
    # Botões de rádio para pagamento
    pagamento_label = tk.Label(verify_window, text="Pagamento:")
    pagamento_label.pack()

    pagamento_value = tk.StringVar()

    pix_radio = tk.Radiobutton(verify_window, text="PIX", variable=pagamento_value, value="PIX")
    pix_radio.pack()

    boleto_radio = tk.Radiobutton(verify_window, text="Boleto", variable=pagamento_value, value="Boleto")
    boleto_radio.pack()

    # Botões de rádio para forma de envio
    entrega_label = tk.Label(verify_window, text="Forma de Envio:")
    entrega_label.pack()

    entrega_value = tk.StringVar()

    email_radio = tk.Radiobutton(verify_window, text="E-mail", variable=entrega_value, value="E-mail")
    email_radio.pack()

    whatsapp_radio = tk.Radiobutton(verify_window, text="WhatsApp", variable=entrega_value, value="WhatsApp")
    whatsapp_radio.pack()
    



def gerar_documento(operador, nome_fantasia, razao_social, cnpj, ramo_atividade, endereco, cidade, cep, telefone, autorizante, cargo, data_venda, vencimento, valor, email_destinatario, alteracoes):
    # Criar um novo documento DOCX
    doc = Document()

    # Adicionar as informações do corpo do email ao documento
    doc.add_heading('ENVIO DE VENDA', level=1)
    doc.add_paragraph('')
    doc.add_paragraph(f'Operador: {operador.upper()}')
    doc.add_paragraph(f'Nome Fantasia: {nome_fantasia.upper()}')
    doc.add_paragraph(f'Razão Social: {razao_social.upper()}')
    doc.add_paragraph(f'CNPJ/CPF: {cnpj.upper()}')
    doc.add_paragraph(f'Ramo de Atividade: {ramo_atividade.upper()}')
    doc.add_paragraph(f'Endereço: {endereco.upper()}')
    doc.add_paragraph(f'Cidade: {cidade.upper()}')
    doc.add_paragraph(f'CEP: {cep.upper()}')
    doc.add_paragraph(f'Telefone: {telefone.upper()}')
    doc.add_paragraph(f'Autorizante: {autorizante.upper()}')
    doc.add_paragraph(f'Cargo: {cargo.upper()}')
    doc.add_paragraph(f'Data de Venda: {data_venda.upper()}')
    doc.add_paragraph(f'Vencimento: {vencimento.upper()}')
    doc.add_paragraph(f'Valor: {valor.upper()}')
    doc.add_paragraph(f'Email: {email_destinatario.upper()}')
    doc.add_paragraph(f'Pagamento selecionado: {pagamento_value.get().upper()}')
    doc.add_paragraph(f'Forma de Envio:')
    doc.add_paragraph('Email' if entrega_value_mail.get() == 1 else '')
    doc.add_paragraph('WhatsApp' if entrega_value_whats.get() == 1 else '')


    doc.add_paragraph('')
    doc.add_paragraph('Alterações:')
    doc.add_paragraph(alteracoes.upper())

    # Salvar o documento em um arquivo DOCX
    docx_file = 'informacoes_email.docx'
    doc.save(docx_file)

    return docx_file

def send_email():
    #mail info
    operador = operador_entry.get()
    nome_fantasia = nome_fantasia_entry.get()
    razao_social = razao_social_entry.get()
    cnpj = cnpj_entry.get()
    ramo_atividade = ramo_atividade_entry.get()
    endereco = endereco_entry.get()
    cidade = cidade_entry.get()
    cep = cep_entry.get()
    telefone = telefone_entry.get()
    autorizante = autorizante_entry.get()
    cargo = cargo_entry.get()
    data_venda = data_venda_entry.get()
    vencimento = vencimento_entry.get()
    valor = valor_entry.get()
    email_destinatario = email_entry.get()
    pagamento_selecionado = pagamento_value.get()
    entrega_email_selecionado = entrega_value_mail.get()
    entrega_whatsapp_selecionado = entrega_value_whats.get()
    alteracoes = alteracoes_text.get("1.0", tk.END)
    
    # Gerar o arquivo DOCX com as informações do corpo do email
    docx_file = gerar_documento(operador, nome_fantasia, razao_social, cnpj, ramo_atividade, endereco, cidade, cep, telefone, autorizante, cargo, data_venda, vencimento, valor, email_destinatario, alteracoes)

    # Criando o objeto MIMEMultipart
    msg = MIMEMultipart()
    
    #body mail
    corpo_email = f"""
    <h1>ENVIO DEVENDA</h1>
    <br>
    <hr>
    <p><strong>Operador:</strong> <u>{operador.upper()}</u></p>
    <p><strong>Nome Fantasia:</strong> <u>{nome_fantasia.upper()}</u></p>
    <p><strong>Razão Social:</strong> <u>{razao_social.upper()}</u></p>
    <p><strong>CNPJ/CPF:</strong> <u>{cnpj.upper()}</u></p>
    <p><strong>Ramo de Atividade:</strong> <u>{ramo_atividade.upper()}</u></p>
    <p><strong>Endereço:</strong> <u>{endereco.upper()}</u></p>
    <p><strong>Cidade:</strong> <u>{cidade.upper()}</u></p>
    <p><strong>CEP:</strong> <u>{cep.upper()}</u></p>
    <p><strong>Telefone:</strong> <u>{telefone.upper()}</u></p>
    <p><strong>Autorizante:</strong> <u>{autorizante.upper()}</u></p>
    <p><strong>Cargo:</strong> <u>{cargo.upper()}</u></p>
    <p><strong>Data de Venda:</strong> <u>{data_venda.upper()}</u></p>
    <p><strong>Vencimento:</strong> <u>{vencimento.upper()}</u></p>
    <p><strong>Valor:</strong> <u>{valor.upper()}</u></p>
    <p><strong>Email:</strong> <u>{email_destinatario.upper()}</u></p>
    <p><strong>Forma de pagamento:</strong> <u>{pagamento_selecionado.upper()}</u></p>
    <p><strong>Forma de Envio:</strong></p>
    <p>{'Email' if entrega_email_selecionado else ''}</p>
    <p>{'WhatsApp' if entrega_whatsapp_selecionado else ''}</p>
    <br>
    <hr>
    <p><strong>Alterações:</strong></p>
    <p>{alteracoes.upper()}</p></u>
    """

    #to send
    msg.attach(MIMEText(corpo_email, 'html'))

    # Adicionando o arquivo DOCX como anexo
    attachment = open(docx_file, 'rb')

    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition', "attachment", filename=os.path.basename(docx_file))

    msg.attach(part)

    # Configurando os dados do email
    msg['Subject'] = 'Assunto'
    msg['From'] = "email que enviará as informações"
    msg['To'] = 'email que receberá as informações'
    password = 'chave de App do email que enviará'

    # Enviando o email
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(msg['From'], password)
    server.send_message(msg)
    server.quit()

    print('Email enviado!')
    
    # Exibir popup de sucesso
    mensagem_sucesso = f'EMAIL ENVIADO COM SUCESSO! Às {datetime.now().strftime("%H:%M:%S")} do dia {datetime.now().strftime("%d/%m/%Y")}'
    messagebox.showinfo("Sucesso", mensagem_sucesso)

    
    nome_fantasia_entry.delete(0, tk.END)
    razao_social_entry.delete(0, tk.END)
    cnpj_entry.delete(0, tk.END)
    ramo_atividade_entry.delete(0, tk.END)
    endereco_entry.delete(0, tk.END)
    cidade_entry.delete(0, tk.END)
    cep_entry.delete(0, tk.END)
    telefone_entry.delete(0, tk.END)
    autorizante_entry.delete(0, tk.END)
    cargo_entry.delete(0, tk.END)
    data_venda_entry.delete(0, tk.END)
    vencimento_entry.delete(0, tk.END)
    valor_entry.delete(0, tk.END)
    email_entry.delete(0, tk.END)
    pagamento_value.set("")
    entrega_value_mail.set(0)  
    entrega_value_whats.set(0)
    alteracoes_text.delete(1.0, tk.END)
    
# create Screen root
root = tk.Tk()
root.title("Program Name")

form_container = tk.Frame(root)
form_container.pack(padx=10, pady=10)

# firt label
operador_label = tk.Label(form_container, text="Field label name:")
operador_label.grid(row=0, column=0, sticky=tk.E)
operador_entry = tk.Entry(form_container)
operador_entry.grid(row=0, column=1)
operador_entry.insert(0, "Inside label text")


# second
nome_fantasia_label = tk.Label(form_container, text="Nome Fantasia:")
nome_fantasia_label.grid(row=1, column=0, sticky=tk.E)
nome_fantasia_entry = tk.Entry(form_container)
nome_fantasia_entry.grid(row=1, column=1)
nome_fantasia_entry.insert(0, "Nome Fantasia")


razao_social_label = tk.Label(form_container, text="Razão Social:")
razao_social_label.grid(row=2, column=0, sticky=tk.E)
razao_social_entry = tk.Entry(form_container)
razao_social_entry.grid(row=2, column=1)
razao_social_entry.insert(0, "Razão Social")


cnpj_label = tk.Label(form_container, text="CNPJ/CPF:")
cnpj_label.grid(row=3, column=0, sticky=tk.E)
cnpj_entry = tk.Entry(form_container)
cnpj_entry.grid(row=3, column=1)
cnpj_entry.insert(0, "CNPJ/CPF")


ramo_atividade_label = tk.Label(form_container, text="Ramo de Atividade:")
ramo_atividade_label.grid(row=4, column=0, sticky=tk.E)
ramo_atividade_entry = tk.Entry(form_container)
ramo_atividade_entry.grid(row=4, column=1)
ramo_atividade_entry.insert(0, "Ramo de Atividade")


endereco_label = tk.Label(form_container, text="Endereço:")
endereco_label.grid(row=5, column=0, sticky=tk.E)
endereco_entry = tk.Entry(form_container)
endereco_entry.grid(row=5, column=1)
endereco_entry.insert(0, "Endereço")


cidade_label = tk.Label(form_container, text="Cidade:")
cidade_label.grid(row=6, column=0, sticky=tk.E)
cidade_entry = tk.Entry(form_container)
cidade_entry.grid(row=6, column=1)
cidade_entry.insert(0, "Cidade")


cep_label = tk.Label(form_container, text="CEP:")
cep_label.grid(row=7, column=0, sticky=tk.E)
cep_entry = tk.Entry(form_container)
cep_entry.grid(row=7, column=1)
cep_entry.insert(0, "CEP")


telefone_label = tk.Label(form_container, text="Telefone:")
telefone_label.grid(row=8, column=0, sticky=tk.E)
telefone_entry = tk.Entry(form_container)
telefone_entry.grid(row=8, column=1)
telefone_entry.insert(0, "Telefone")


autorizante_label = tk.Label(form_container, text="Autorizante:")
autorizante_label.grid(row=9, column=0, sticky=tk.E)
autorizante_entry = tk.Entry(form_container)
autorizante_entry.grid(row=9, column=1)
autorizante_entry.insert(0, "Autorizante")


cargo_label = tk.Label(form_container, text="Cargo:")
cargo_label.grid(row=10, column=0, sticky=tk.E)
cargo_entry = tk.Entry(form_container)
cargo_entry.grid(row=10, column=1)
cargo_entry.insert(0, "Cargo")


data_venda_label = tk.Label(form_container, text="Data Venda:")
data_venda_label.grid(row=11, column=0, sticky=tk.E)
data_venda_entry = tk.Entry(form_container)
data_venda_entry.grid(row=11, column=1)
data_venda_entry.insert(0, "Data Venda")


vencimento_label = tk.Label(form_container, text="Vencimento:")
vencimento_label.grid(row=12, column=0, sticky=tk.E)
vencimento_entry = tk.Entry(form_container)
vencimento_entry.grid(row=12, column=1)
vencimento_entry.insert(0, "Vencimento")


valor_label = tk.Label(form_container, text="Valor:")
valor_label.grid(row=13, column=0, sticky=tk.E)
valor_entry = tk.Entry(form_container)
valor_entry.grid(row=13, column=1)
valor_entry.insert(0, "Valor")


email_label = tk.Label(form_container, text="Email:")
email_label.grid(row=14, column=0, sticky=tk.E)
email_entry = tk.Entry(form_container)
email_entry.grid(row=14, column=1)
email_entry.insert(0, "Email")

# Botões de rádio para pagamento


pagamento_value = tk.StringVar()
pix_radio = tk.Radiobutton(form_container, text="PIX", variable=pagamento_value, value="PIX")
pix_radio.grid(row=15, column=0, sticky=tk.E)
boleto_radio = tk.Radiobutton(form_container, text="Boleto", variable=pagamento_value, value="Boleto")
boleto_radio.grid(row=16, column=0, sticky=tk.E)


# Botões de rádio para forma de envio

entrega_value_mail = tk.IntVar()
entrega_value_whats = tk.IntVar()
email_radio = tk.Checkbutton(form_container, text="E-mail", variable=entrega_value_mail, onvalue=1, offvalue=0)
email_radio.grid(row=17, column=0, sticky=tk.E)
whatsapp_radio = tk.Checkbutton(form_container, text="WhatsApp", variable=entrega_value_whats, onvalue=1, offvalue=0)
whatsapp_radio.grid(row=18, column=0, sticky=tk.E)


alteracoes_label = tk.Label(form_container, text="Alterações:")
alteracoes_label.grid(row=20, column=0, sticky=tk.E)
alteracoes_text = tk.Text(form_container, width=30, height=10)
alteracoes_text.grid(row=20, column=1, sticky=tk.NSEW)
alteracoes_text.insert(tk.END, "Insira suas alterações aqui")



verificar_button = tk.Button(root, text="Verificar informações", command=verify)
verificar_button.pack(pady=10)
enviar_button = tk.Button(root, text="Enviar", command=send_email)
enviar_button.pack(pady=10)

root.mainloop()
